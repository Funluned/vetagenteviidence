from __future__ import annotations

import csv
import hashlib
from io import BytesIO, StringIO

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import Workbook, load_workbook

from vetevidence.mechanism_prediction import SourceProvenance
from vetevidence.network_files import (
    COMPOUND_TARGET_COLUMNS,
    MAX_NETWORK_DATA_ROWS,
    TARGET_PATHWAY_COLUMNS,
    analyze_network_pharmacology_files,
    compound_target_template_docx,
    compound_target_template_xlsx,
    network_result_to_docx,
    network_result_to_xlsx,
    parse_compound_target_file,
    parse_target_pathway_file,
    target_pathway_template_docx,
    target_pathway_template_xlsx,
)


COMPOUND_TARGET_ROWS = [
    [
        "Compound A",
        "CID:1",
        "Target bacterium",
        "Target One",
        "UniProt:P11111",
    ],
    [
        "Compound B",
        "CID:2",
        "Target bacterium",
        "Target One",
        "UniProt:P11111",
    ],
    [
        "Compound A",
        "CID:1",
        "Target bacterium",
        "Target Two",
        "UniProt:P22222",
    ],
]
TARGET_PATHWAY_ROWS = [
    [
        "Target bacterium",
        "Target One",
        "UniProt:P11111",
        "Pathway A",
        "KEGG:map00010",
    ],
    [
        "Target bacterium",
        "Target One",
        "UniProt:P11111",
        "Pathway B",
        "KEGG:map00020",
    ],
    [
        "Target bacterium",
        "Target Two",
        "UniProt:P22222",
        "Pathway C",
        "KEGG:map00030",
    ],
]


def source(name: str, accession: str) -> SourceProvenance:
    return SourceProvenance(
        source_name=name,
        accession=accession,
        version="2026-07-30",
    )


def csv_payload(headers: tuple[str, ...], rows: list[list[str]]) -> bytes:
    output = StringIO(newline="")
    writer = csv.writer(output, lineterminator="\n")
    writer.writerow(headers)
    writer.writerows(rows)
    return output.getvalue().encode("utf-8")


def xlsx_payload(headers: tuple[str, ...], rows: list[list[str]]) -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(list(headers))
    for row in rows:
        sheet.append(row)
    output = BytesIO()
    workbook.save(output)
    workbook.close()
    return output.getvalue()


def docx_payload(headers: tuple[str, ...], rows: list[list[str]]) -> bytes:
    document = Document()
    table = document.add_table(rows=1, cols=len(headers))
    for cell, value in zip(table.rows[0].cells, headers, strict=True):
        cell.text = value
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values, strict=True):
            cell.text = value
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def payload_for(
    suffix: str,
    headers: tuple[str, ...],
    rows: list[list[str]],
) -> bytes:
    if suffix == ".csv":
        return csv_payload(headers, rows)
    if suffix == ".xlsx":
        return xlsx_payload(headers, rows)
    if suffix == ".docx":
        return docx_payload(headers, rows)
    raise AssertionError(f"unexpected test suffix: {suffix}")


def analyze(suffix: str, *, compound_rows=None, pathway_rows=None):
    compound_payload = payload_for(
        suffix,
        COMPOUND_TARGET_COLUMNS,
        compound_rows or COMPOUND_TARGET_ROWS,
    )
    pathway_payload = payload_for(
        suffix,
        TARGET_PATHWAY_COLUMNS,
        pathway_rows or TARGET_PATHWAY_ROWS,
    )
    result = analyze_network_pharmacology_files(
        compound_payload,
        pathway_payload,
        compound_target_filename=f"compound-target{suffix}",
        target_pathway_filename=f"target-pathway{suffix}",
        compound_target_source=source("compound-target", "dataset:ct"),
        target_pathway_source=source("target-pathway", "dataset:tp"),
    )
    return result, compound_payload, pathway_payload


@pytest.mark.parametrize("suffix", [".csv", ".xlsx", ".docx"])
def test_csv_xlsx_docx_inputs_are_semantically_equivalent(suffix: str) -> None:
    expected, _, _ = analyze(".csv")
    actual, compound_payload, pathway_payload = analyze(suffix)

    assert actual.model_dump(exclude={"sources"}) == expected.model_dump(
        exclude={"sources"}
    )
    assert actual.sources[0].sha256 == hashlib.sha256(
        compound_payload
    ).hexdigest()
    assert actual.sources[1].sha256 == hashlib.sha256(
        pathway_payload
    ).hexdigest()
    assert actual.ranked_targets[0].source_rows[0].row_number == 2


@pytest.mark.parametrize(
    ("template_factory", "headers", "row", "filename", "expected"),
    [
        (
            compound_target_template_xlsx,
            COMPOUND_TARGET_COLUMNS,
            COMPOUND_TARGET_ROWS[0],
            "compound-target.xlsx",
            ("Compound A", "UniProt:P11111"),
        ),
        (
            target_pathway_template_xlsx,
            TARGET_PATHWAY_COLUMNS,
            TARGET_PATHWAY_ROWS[0],
            "target-pathway.xlsx",
            ("Pathway A", "KEGG:map00010"),
        ),
        (
            compound_target_template_docx,
            COMPOUND_TARGET_COLUMNS,
            COMPOUND_TARGET_ROWS[0],
            "compound-target.docx",
            ("Compound A", "UniProt:P11111"),
        ),
        (
            target_pathway_template_docx,
            TARGET_PATHWAY_COLUMNS,
            TARGET_PATHWAY_ROWS[0],
            "target-pathway.docx",
            ("Pathway A", "KEGG:map00010"),
        ),
    ],
)
def test_generated_templates_can_be_filled_and_parsed(
    template_factory,
    headers: tuple[str, ...],
    row: list[str],
    filename: str,
    expected: tuple[str, str],
) -> None:
    payload = template_factory()
    if filename.endswith(".xlsx"):
        workbook = load_workbook(BytesIO(payload))
        sheet = workbook.active
        assert tuple(cell.value for cell in sheet[1]) == headers
        for column, value in enumerate(row, start=1):
            sheet.cell(2, column, value)
        output = BytesIO()
        workbook.save(output)
        workbook.close()
        filled = output.getvalue()
    else:
        document = Document(BytesIO(payload))
        assert tuple(
            cell.text for cell in document.tables[0].rows[0].cells
        ) == headers
        for cell, value in zip(
            document.tables[0].rows[1].cells,
            row,
            strict=True,
        ):
            cell.text = value
        output = BytesIO()
        document.save(output)
        filled = output.getvalue()

    if filename.startswith("compound"):
        records = parse_compound_target_file(
            filled,
            filename=filename,
            source=source(filename, "dataset:template"),
        )
        actual = (records[0].compound, records[0].target_accession)
    else:
        records = parse_target_pathway_file(
            filled,
            filename=filename,
            source=source(filename, "dataset:template"),
        )
        actual = (records[0].pathway, records[0].pathway_accession)

    assert actual == expected
    assert records[0].row_number == 2
    assert records[0].source.sha256 == hashlib.sha256(filled).hexdigest()


def test_xlsx_rejects_formula_cells() -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(COMPOUND_TARGET_COLUMNS)
    sheet.append(COMPOUND_TARGET_ROWS[0])
    sheet["A2"] = "=1+1"
    output = BytesIO()
    workbook.save(output)
    workbook.close()

    with pytest.raises(ValueError, match="公式"):
        parse_compound_target_file(
            output.getvalue(),
            filename="formula.xlsx",
            source=source("formula.xlsx", "dataset:bad"),
        )


@pytest.mark.parametrize("merge_range", ["A1:B1", "A2:A3"])
def test_xlsx_rejects_merged_cells(merge_range: str) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(COMPOUND_TARGET_COLUMNS)
    sheet.append(COMPOUND_TARGET_ROWS[0])
    sheet.append(COMPOUND_TARGET_ROWS[1])
    sheet.merge_cells(merge_range)
    output = BytesIO()
    workbook.save(output)
    workbook.close()

    with pytest.raises(ValueError, match="合并单元格"):
        parse_compound_target_file(
            output.getvalue(),
            filename="merged.xlsx",
            source=source("merged.xlsx", "dataset:bad"),
        )


def test_xlsx_rejects_multiple_nonempty_worksheets() -> None:
    workbook = Workbook()
    workbook.active.append(COMPOUND_TARGET_COLUMNS)
    workbook.active.append(COMPOUND_TARGET_ROWS[0])
    workbook.create_sheet("another").append(["unexpected"])
    output = BytesIO()
    workbook.save(output)
    workbook.close()

    with pytest.raises(ValueError, match="一个非空工作表"):
        parse_compound_target_file(
            output.getvalue(),
            filename="multiple.xlsx",
            source=source("multiple.xlsx", "dataset:bad"),
        )


@pytest.mark.parametrize("suffix", [".csv", ".xlsx", ".docx"])
def test_network_inputs_require_template_headers_in_exact_order(
    suffix: str,
) -> None:
    reordered = (
        COMPOUND_TARGET_COLUMNS[1],
        COMPOUND_TARGET_COLUMNS[0],
        *COMPOUND_TARGET_COLUMNS[2:],
    )
    payload = payload_for(suffix, reordered, COMPOUND_TARGET_ROWS)

    with pytest.raises(ValueError, match="表头|列顺序"):
        parse_compound_target_file(
            payload,
            filename=f"reordered{suffix}",
            source=source(f"reordered{suffix}", "dataset:bad"),
        )


@pytest.mark.parametrize("suffix", [".csv", ".xlsx", ".docx"])
def test_network_inputs_reject_extra_columns(suffix: str) -> None:
    headers = (*COMPOUND_TARGET_COLUMNS, "unapproved_note")
    rows = [[*row, "not part of the template"] for row in COMPOUND_TARGET_ROWS]
    payload = payload_for(suffix, headers, rows)

    with pytest.raises(ValueError, match="表头|额外"):
        parse_compound_target_file(
            payload,
            filename=f"extra{suffix}",
            source=source(f"extra{suffix}", "dataset:bad"),
        )


@pytest.mark.parametrize("suffix", [".csv", ".xlsx", ".docx"])
def test_network_inputs_reject_missing_columns(suffix: str) -> None:
    headers = COMPOUND_TARGET_COLUMNS[:-1]
    rows = [row[:-1] for row in COMPOUND_TARGET_ROWS]
    payload = payload_for(suffix, headers, rows)

    with pytest.raises(ValueError, match="target_accession"):
        parse_compound_target_file(
            payload,
            filename=f"missing{suffix}",
            source=source(f"missing{suffix}", "dataset:bad"),
        )


@pytest.mark.parametrize("suffix", [".csv", ".xlsx", ".docx"])
def test_network_inputs_reject_duplicate_columns(suffix: str) -> None:
    headers = (*COMPOUND_TARGET_COLUMNS, "target_accession")
    rows = [[*row, "UniProt:WRONG"] for row in COMPOUND_TARGET_ROWS]
    payload = payload_for(suffix, headers, rows)

    with pytest.raises(ValueError, match="重复列"):
        parse_compound_target_file(
            payload,
            filename=f"duplicate{suffix}",
            source=source(f"duplicate{suffix}", "dataset:bad"),
        )


def test_csv_rejects_data_rows_wider_than_the_header() -> None:
    payload = csv_payload(
        COMPOUND_TARGET_COLUMNS,
        [[*COMPOUND_TARGET_ROWS[0], "silently shifted value"]],
    )

    with pytest.raises(ValueError, match="列数多于表头"):
        parse_compound_target_file(
            payload,
            filename="wide-row.csv",
            source=source("wide-row.csv", "dataset:bad"),
        )


def test_docx_rejects_multiple_nonempty_tables() -> None:
    document = Document()
    for _ in range(2):
        table = document.add_table(rows=2, cols=len(COMPOUND_TARGET_COLUMNS))
        for cell, value in zip(
            table.rows[0].cells,
            COMPOUND_TARGET_COLUMNS,
            strict=True,
        ):
            cell.text = value
        for cell, value in zip(
            table.rows[1].cells,
            COMPOUND_TARGET_ROWS[0],
            strict=True,
        ):
            cell.text = value
    output = BytesIO()
    document.save(output)

    with pytest.raises(ValueError, match="一个非空表格"):
        parse_compound_target_file(
            output.getvalue(),
            filename="multiple.docx",
            source=source("multiple.docx", "dataset:bad"),
        )


def _docx_with_merge_marker(marker: str) -> bytes:
    document = Document(
        BytesIO(
            docx_payload(
                COMPOUND_TARGET_COLUMNS,
                COMPOUND_TARGET_ROWS,
            )
        )
    )
    cell_properties = document.tables[0].cell(1, 0)._tc.get_or_add_tcPr()
    element = OxmlElement(f"w:{marker}")
    element.set(qn("w:val"), "restart")
    cell_properties.append(element)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


@pytest.mark.parametrize("marker", ["gridSpan", "vMerge", "hMerge"])
def test_docx_rejects_all_merge_encodings(marker: str) -> None:
    with pytest.raises(ValueError, match="合并单元格"):
        parse_compound_target_file(
            _docx_with_merge_marker(marker),
            filename=f"{marker}.docx",
            source=source(f"{marker}.docx", "dataset:bad"),
        )


def test_docx_rejects_nested_tables() -> None:
    document = Document(
        BytesIO(
            docx_payload(
                COMPOUND_TARGET_COLUMNS,
                COMPOUND_TARGET_ROWS,
            )
        )
    )
    nested = document.tables[0].cell(1, 0).add_table(rows=1, cols=1)
    nested.cell(0, 0).text = "hidden nested data"
    output = BytesIO()
    document.save(output)

    with pytest.raises(ValueError, match="嵌套表格"):
        parse_compound_target_file(
            output.getvalue(),
            filename="nested.docx",
            source=source("nested.docx", "dataset:bad"),
        )


def test_csv_rejects_more_than_ten_thousand_nonempty_rows() -> None:
    row = COMPOUND_TARGET_ROWS[0]
    payload = csv_payload(
        COMPOUND_TARGET_COLUMNS,
        [row] * (MAX_NETWORK_DATA_ROWS + 1),
    )

    with pytest.raises(ValueError, match="10,000"):
        parse_compound_target_file(
            payload,
            filename="too-many.csv",
            source=source("too-many.csv", "dataset:bad"),
        )


def injection_result():
    compound_rows = [
        [
            '=HYPERLINK("https://invalid.example")',
            "CID:1",
            "Target bacterium",
            "Target One",
            "UniProt:P11111",
        ]
    ]
    pathway_rows = [
        [
            "Target bacterium",
            "Target One",
            "UniProt:P11111",
            "@SUM(1,1)",
            "KEGG:map00010",
        ]
    ]
    result, _, _ = analyze(
        ".csv",
        compound_rows=compound_rows,
        pathway_rows=pathway_rows,
    )
    return result


def test_xlsx_result_export_is_traceable_and_formula_safe() -> None:
    payload = network_result_to_xlsx(injection_result())
    workbook = load_workbook(BytesIO(payload), data_only=False)

    assert workbook.sheetnames == ["摘要", "靶点排名", "靶点-通路", "来源"]
    assert workbook["摘要"]["A1"].value == "项目"
    assert workbook["靶点排名"]["E2"].value.startswith("=HYPERLINK")
    assert workbook["靶点排名"]["E2"].data_type == "s"
    assert workbook["靶点排名"]["E2"].quotePrefix is True
    assert workbook["靶点-通路"]["E2"].value == "@SUM(1,1)"
    assert workbook["靶点-通路"]["E2"].data_type == "s"
    assert all(
        cell.data_type != "f"
        for sheet in workbook.worksheets
        for row in sheet.iter_rows()
        for cell in row
    )
    assert workbook["来源"]["D2"].value == injection_result().sources[0].sha256
    workbook.close()


def test_docx_result_export_is_traceable_and_has_no_word_fields() -> None:
    payload = network_result_to_docx(injection_result())
    document = Document(BytesIO(payload))

    assert document.paragraphs[0].text == (
        "VetEvidence 网络药理学靶点与通路报告"
    )
    assert document.paragraphs[0].text.strip()
    assert len(document.tables) == 4
    assert document.tables[1].cell(1, 4).text.startswith("=HYPERLINK")
    assert document.tables[2].cell(1, 2).text == "@SUM(1,1)"
    assert document.tables[3].cell(1, 3).text == (
        injection_result().sources[0].sha256
    )
    assert not document.element.xpath(".//w:fldSimple | .//w:instrText")


def test_docx_result_tables_fit_inside_page_content_width() -> None:
    document = Document(BytesIO(network_result_to_docx(injection_result())))
    section = document.sections[0]
    content_width = (
        section.page_width.twips
        - section.left_margin.twips
        - section.right_margin.twips
    )

    for table in document.tables:
        table_width = int(table._tbl.tblPr.xpath("./w:tblW")[0].get(qn("w:w")))
        indent_nodes = table._tbl.tblPr.xpath("./w:tblInd")
        indent = (
            int(indent_nodes[0].get(qn("w:w")))
            if indent_nodes
            else 0
        )
        grid_width = sum(
            int(column.get(qn("w:w")))
            for column in table._tbl.tblGrid.gridCol_lst
        )
        assert table_width == grid_width
        assert max(indent, 0) + table_width <= content_width
