"""Tabular file adapters and exports for network pharmacology."""

from __future__ import annotations

import csv
import hashlib
from datetime import date, datetime
from io import BytesIO, StringIO
from pathlib import Path
from zipfile import BadZipFile, ZipFile, is_zipfile

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook, load_workbook
from openpyxl.cell.cell import Cell
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

from vetevidence.mechanism_prediction import (
    CompoundTargetRecord,
    NetworkPharmacologyParameters,
    NetworkPharmacologyResult,
    SourceProvenance,
    TargetPathwayRecord,
    analyze_network_pharmacology_records,
    parse_compound_target_csv,
    parse_target_pathway_csv,
)


COMPOUND_TARGET_COLUMNS = (
    "compound",
    "compound_accession",
    "organism",
    "target",
    "target_accession",
)
TARGET_PATHWAY_COLUMNS = (
    "organism",
    "target",
    "target_accession",
    "pathway",
    "pathway_accession",
)
SUPPORTED_NETWORK_SUFFIXES = {".csv", ".xlsx", ".docx"}
MAX_NETWORK_FILE_BYTES = 10 * 1024 * 1024
MAX_NETWORK_DATA_ROWS = 10_000
MAX_NETWORK_COLUMNS = 64
MAX_ZIP_MEMBERS = 2_000
MAX_ZIP_EXPANDED_BYTES = 64 * 1024 * 1024

_BLUE = "1F4E78"
_DARK_BLUE = "17365D"
_LIGHT_BLUE = "D9EAF7"
_LIGHT_GRAY = "F2F4F7"
_WHITE = "FFFFFF"
_THIN_GRAY = Side(style="thin", color="D9E1F2")


def _checked_payload(payload: bytes, *, source_name: str) -> bytes:
    if not payload:
        raise ValueError(f"{source_name} 为空。")
    if len(payload) > MAX_NETWORK_FILE_BYTES:
        raise ValueError(f"{source_name} 超过 10 MB，拒绝解析。")
    return payload


def _source_with_digest(
    source: SourceProvenance,
    payload: bytes,
) -> SourceProvenance:
    digest = hashlib.sha256(payload).hexdigest()
    if source.sha256 is not None and source.sha256.casefold() != digest:
        raise ValueError(f"{source.source_name} 的 SHA-256 与原文件不一致。")
    return source.model_copy(update={"sha256": digest})


def _inspect_ooxml_archive(payload: bytes, *, source_name: str) -> None:
    if not is_zipfile(BytesIO(payload)):
        raise ValueError(f"{source_name} 不是有效的 OOXML 文件。")
    try:
        with ZipFile(BytesIO(payload)) as archive:
            members = archive.infolist()
    except BadZipFile as exc:
        raise ValueError(f"{source_name} 不是有效的 OOXML 文件。") from exc
    if len(members) > MAX_ZIP_MEMBERS:
        raise ValueError(f"{source_name} 包含过多压缩成员，拒绝解析。")
    expanded = sum(member.file_size for member in members)
    if expanded > MAX_ZIP_EXPANDED_BYTES:
        raise ValueError(f"{source_name} 解压后超过 64 MB，拒绝解析。")


def _text_value(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.isoformat()
    if isinstance(value, date):
        return value.isoformat()
    if isinstance(value, float):
        return format(value, ".15g")
    return str(value).strip()


def _validated_headers(
    values: list[str],
    *,
    required_columns: tuple[str, ...],
    source_name: str,
) -> list[str]:
    if not values or not any(values):
        raise ValueError(f"{source_name} 第一行必须是字段表头。")
    if any(not value for value in values):
        raise ValueError(f"{source_name} 表头包含空列名。")
    duplicates = sorted(
        {value for value in values if values.count(value) > 1}
    )
    if duplicates:
        raise ValueError(
            f"{source_name} 表头包含重复列：" + "、".join(duplicates)
        )
    missing = sorted(set(required_columns) - set(values))
    if missing:
        raise ValueError(
            f"{source_name} 缺少必需列：" + "、".join(missing)
        )
    if tuple(values) != required_columns:
        raise ValueError(
            f"{source_name} 表头必须与模板完全一致并保持顺序："
            + "、".join(required_columns)
        )
    return values


def _rows_from_xlsx(
    payload: bytes,
    *,
    required_columns: tuple[str, ...],
    source: SourceProvenance,
) -> tuple[list[tuple[int, dict[str, str]]], SourceProvenance]:
    raw = _checked_payload(payload, source_name=source.source_name)
    _inspect_ooxml_archive(raw, source_name=source.source_name)
    try:
        workbook = load_workbook(
            BytesIO(raw),
            read_only=False,
            data_only=False,
            keep_links=False,
        )
    except Exception as exc:
        raise ValueError(f"{source.source_name} 不是可读取的 XLSX。") from exc
    try:
        nonempty = []
        for sheet in workbook.worksheets:
            if (
                sheet.max_row > MAX_NETWORK_DATA_ROWS + 1
                or sheet.max_column > MAX_NETWORK_COLUMNS
            ):
                raise ValueError(
                    f"{source.source_name} 工作表 {sheet.title} 超过"
                    " 10,000 行或 64 列。"
                )
            if any(
                _text_value(cell.value)
                for row in sheet.iter_rows()
                for cell in row
            ):
                nonempty.append(sheet)
        if len(nonempty) != 1:
            raise ValueError(
                f"{source.source_name} 必须且只能包含一个非空工作表；"
                f"实际为 {len(nonempty)} 个。"
            )
        sheet = nonempty[0]
        if sheet.merged_cells.ranges:
            raise ValueError(
                f"{source.source_name} 不支持合并单元格。"
            )
        for row in sheet.iter_rows():
            for cell in row:
                if cell.data_type == "f":
                    raise ValueError(
                        f"{source.source_name} 第 {cell.row} 行包含公式，"
                        "请改为静态值。"
                    )
        headers = _validated_headers(
            [_text_value(cell.value) for cell in sheet[1]],
            required_columns=required_columns,
            source_name=source.source_name,
        )
        rows: list[tuple[int, dict[str, str]]] = []
        for row_number, cells in enumerate(
            sheet.iter_rows(min_row=2, max_col=len(headers)),
            start=2,
        ):
            values = [_text_value(cell.value) for cell in cells]
            if not any(values):
                continue
            row = dict(zip(headers, values, strict=True))
            empty = sorted(
                column for column in required_columns if not row.get(column)
            )
            if empty:
                raise ValueError(
                    f"{source.source_name} 第 {row_number} 行缺少值："
                    + "、".join(empty)
                )
            rows.append((row_number, row))
        if not rows:
            raise ValueError(f"{source.source_name} 没有可用数据行。")
        return rows, _source_with_digest(source, raw)
    finally:
        workbook.close()


def _rows_from_docx(
    payload: bytes,
    *,
    required_columns: tuple[str, ...],
    source: SourceProvenance,
) -> tuple[list[tuple[int, dict[str, str]]], SourceProvenance]:
    raw = _checked_payload(payload, source_name=source.source_name)
    _inspect_ooxml_archive(raw, source_name=source.source_name)
    try:
        document = Document(BytesIO(raw))
    except Exception as exc:
        raise ValueError(f"{source.source_name} 不是可读取的 DOCX。") from exc
    nonempty_tables = [
        table
        for table in document.tables
        if any(cell.text.strip() for row in table.rows for cell in row.cells)
    ]
    if len(nonempty_tables) != 1:
        raise ValueError(
            f"{source.source_name} 必须且只能包含一个非空表格；"
            f"实际为 {len(nonempty_tables)} 个。"
        )
    table = nonempty_tables[0]
    if len(table.rows) > MAX_NETWORK_DATA_ROWS + 1:
        raise ValueError(f"{source.source_name} 表格超过 10,000 行。")
    if not table.rows or len(table.rows[0].cells) > MAX_NETWORK_COLUMNS:
        raise ValueError(f"{source.source_name} 表格为空或超过 64 列。")
    if table._tbl.xpath(".//w:gridSpan | .//w:vMerge"):
        raise ValueError(f"{source.source_name} 不支持合并单元格。")
    if table._tbl.xpath(".//w:hMerge"):
        raise ValueError(f"{source.source_name} 不支持合并单元格。")
    if table._tbl.xpath(".//w:tbl"):
        raise ValueError(f"{source.source_name} 不支持嵌套表格。")
    if table._tbl.xpath(".//w:fldSimple | .//w:instrText"):
        raise ValueError(
            f"{source.source_name} 表格包含 Word 字段，请改为静态值。"
        )
    headers = _validated_headers(
        [cell.text.strip() for cell in table.rows[0].cells],
        required_columns=required_columns,
        source_name=source.source_name,
    )
    rows: list[tuple[int, dict[str, str]]] = []
    for row_number, table_row in enumerate(table.rows[1:], start=2):
        values = [cell.text.strip() for cell in table_row.cells]
        if len(values) != len(headers):
            raise ValueError(
                f"{source.source_name} 第 {row_number} 行列数与表头不一致。"
            )
        if not any(values):
            continue
        row = dict(zip(headers, values, strict=True))
        empty = sorted(
            column for column in required_columns if not row.get(column)
        )
        if empty:
            raise ValueError(
                f"{source.source_name} 第 {row_number} 行缺少值："
                + "、".join(empty)
            )
        rows.append((row_number, row))
    if not rows:
        raise ValueError(f"{source.source_name} 没有可用数据行。")
    return rows, _source_with_digest(source, raw)


def _file_rows(
    payload: bytes,
    *,
    filename: str,
    required_columns: tuple[str, ...],
    source: SourceProvenance,
) -> tuple[list[tuple[int, dict[str, str]]], SourceProvenance]:
    suffix = Path(filename).suffix.casefold()
    if suffix not in SUPPORTED_NETWORK_SUFFIXES:
        raise ValueError(
            f"{source.source_name} 不支持 {suffix or '无扩展名'}；"
            "请上传 CSV、XLSX 或 DOCX。"
        )
    if suffix == ".xlsx":
        return _rows_from_xlsx(
            payload,
            required_columns=required_columns,
            source=source,
        )
    if suffix == ".docx":
        return _rows_from_docx(
            payload,
            required_columns=required_columns,
            source=source,
        )
    raise AssertionError("CSV 由现有严格解析器处理。")


def _guard_csv_row_count(payload: bytes, *, source_name: str) -> None:
    try:
        text = payload.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise ValueError(f"{source_name} 必须是 UTF-8 CSV。") from exc
    try:
        rows = csv.reader(StringIO(text), strict=True)
        next(rows, None)
        populated_rows = 0
        for row in rows:
            if not any(value.strip() for value in row):
                continue
            populated_rows += 1
            if populated_rows > MAX_NETWORK_DATA_ROWS:
                raise ValueError(f"{source_name} 超过 10,000 个数据行。")
    except csv.Error as exc:
        raise ValueError(f"{source_name} 不是结构完整的 CSV：{exc}") from exc


def parse_compound_target_file(
    payload: bytes,
    *,
    filename: str,
    source: SourceProvenance,
) -> list[CompoundTargetRecord]:
    """Parse one compound-target relation file without losing file provenance."""

    if Path(filename).suffix.casefold() == ".csv":
        raw = _checked_payload(payload, source_name=source.source_name)
        _guard_csv_row_count(raw, source_name=source.source_name)
        records = parse_compound_target_csv(raw, source=source)
    else:
        rows, traced_source = _file_rows(
            payload,
            filename=filename,
            required_columns=COMPOUND_TARGET_COLUMNS,
            source=source,
        )
        records = [
            CompoundTargetRecord(
                compound=row["compound"],
                compound_accession=row["compound_accession"],
                organism=row["organism"],
                target=row["target"],
                target_accession=row["target_accession"],
                source=traced_source,
                row_number=row_number,
            )
            for row_number, row in rows
        ]
    if len(records) > MAX_NETWORK_DATA_ROWS:
        raise ValueError(f"{source.source_name} 超过 10,000 个数据行。")
    return records


def parse_target_pathway_file(
    payload: bytes,
    *,
    filename: str,
    source: SourceProvenance,
) -> list[TargetPathwayRecord]:
    """Parse one target-pathway relation file without losing file provenance."""

    if Path(filename).suffix.casefold() == ".csv":
        raw = _checked_payload(payload, source_name=source.source_name)
        _guard_csv_row_count(raw, source_name=source.source_name)
        records = parse_target_pathway_csv(raw, source=source)
    else:
        rows, traced_source = _file_rows(
            payload,
            filename=filename,
            required_columns=TARGET_PATHWAY_COLUMNS,
            source=source,
        )
        records = [
            TargetPathwayRecord(
                organism=row["organism"],
                target=row["target"],
                target_accession=row["target_accession"],
                pathway=row["pathway"],
                pathway_accession=row["pathway_accession"],
                source=traced_source,
                row_number=row_number,
            )
            for row_number, row in rows
        ]
    if len(records) > MAX_NETWORK_DATA_ROWS:
        raise ValueError(f"{source.source_name} 超过 10,000 个数据行。")
    return records


def analyze_network_pharmacology_files(
    compound_target_payload: bytes,
    target_pathway_payload: bytes,
    *,
    compound_target_filename: str,
    target_pathway_filename: str,
    compound_target_source: SourceProvenance,
    target_pathway_source: SourceProvenance,
    parameters: NetworkPharmacologyParameters | None = None,
) -> NetworkPharmacologyResult:
    compound_records = parse_compound_target_file(
        compound_target_payload,
        filename=compound_target_filename,
        source=compound_target_source,
    )
    pathway_records = parse_target_pathway_file(
        target_pathway_payload,
        filename=target_pathway_filename,
        source=target_pathway_source,
    )
    return analyze_network_pharmacology_records(
        compound_records,
        pathway_records,
        parameters=parameters,
    )


def _excel_value(cell: Cell, value: object) -> None:
    if isinstance(value, str) and value.startswith(("=", "+", "-", "@")):
        cell.value = value
        cell.data_type = "s"
        cell.quotePrefix = True
    else:
        cell.value = value


def _style_sheet(sheet, *, header_row: int = 1) -> None:
    sheet.sheet_view.showGridLines = False
    sheet.freeze_panes = f"A{header_row + 1}"
    sheet.auto_filter.ref = sheet.dimensions
    for cell in sheet[header_row]:
        cell.fill = PatternFill("solid", fgColor=_BLUE)
        cell.font = Font(color=_WHITE, bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = Border(bottom=_THIN_GRAY)
    sheet.row_dimensions[header_row].height = 24
    for column_index in range(1, sheet.max_column + 1):
        values = [
            len(str(sheet.cell(row, column_index).value or ""))
            for row in range(1, sheet.max_row + 1)
        ]
        width = min(max(max(values, default=8) + 2, 10), 42)
        sheet.column_dimensions[get_column_letter(column_index)].width = width
        for row in range(header_row + 1, sheet.max_row + 1):
            sheet.cell(row, column_index).alignment = Alignment(
                vertical="top",
                wrap_text=True,
            )


def _append_excel_rows(sheet, rows: list[list[object]]) -> None:
    row_number = (
        1
        if sheet.max_row == 1
        and sheet.max_column == 1
        and sheet["A1"].value is None
        else sheet.max_row + 1
    )
    for values in rows:
        for column_number, value in enumerate(values, start=1):
            _excel_value(sheet.cell(row_number, column_number), value)
        row_number += 1


def _pathway_rows(result: NetworkPharmacologyResult) -> list[list[object]]:
    rows: list[list[object]] = []
    for target in result.ranked_targets:
        if target.pathways:
            links = [
                (link.pathway, link.pathway_accession)
                for link in target.pathways
            ]
        else:
            links = [("", accession) for accession in target.pathway_accessions]
        source_rows = "；".join(
            f"{reference.source_accession}#row={reference.row_number}"
            for reference in target.source_rows
        )
        for pathway, accession in links:
            rows.append(
                [
                    target.rank,
                    target.organism,
                    target.target,
                    target.target_accession,
                    pathway,
                    accession,
                    source_rows,
                ]
            )
    return rows


def network_result_to_xlsx(result: NetworkPharmacologyResult) -> bytes:
    """Export a formula-free, traceable workbook for targets and pathways."""

    workbook = Workbook()
    summary = workbook.active
    summary.title = "摘要"
    _append_excel_rows(
        summary,
        [
            ["项目", "值"],
            ["证据等级", result.evidence_grade.value],
            ["算法版本", result.parameters.algorithm_version],
            ["排名规则", result.parameters.ranking_method],
            ["输入化合物数", result.summary.input_compound_count],
            ["交集靶点数", result.summary.intersection_target_count],
            ["交集通路数", result.summary.intersection_pathway_count],
            [
                "科研边界",
                "网络排名仅反映用户导入关系的透明拓扑统计，"
                "不能证明靶点有效、药物结合或协同作用。",
            ],
        ],
    )
    _style_sheet(summary)

    targets = workbook.create_sheet("靶点排名")
    _append_excel_rows(
        targets,
        [
            [
                "排名",
                "研究对象",
                "靶点",
                "靶点 accession",
                "化合物",
                "化合物 accession",
                "化合物数",
                "通路数",
                "网络分数",
                "来源行",
            ],
            *[
                [
                    target.rank,
                    target.organism,
                    target.target,
                    target.target_accession,
                    "；".join(link.compound for link in target.compounds),
                    "；".join(target.compound_accessions),
                    target.compound_degree,
                    target.pathway_degree,
                    target.network_score,
                    "；".join(
                        f"{reference.source_accession}#row="
                        f"{reference.row_number}"
                        for reference in target.source_rows
                    ),
                ]
                for target in result.ranked_targets
            ],
        ],
    )
    _style_sheet(targets)

    pathways = workbook.create_sheet("靶点-通路")
    _append_excel_rows(
        pathways,
        [
            [
                "靶点排名",
                "研究对象",
                "靶点",
                "靶点 accession",
                "通路",
                "通路 accession",
                "来源行",
            ],
            *_pathway_rows(result),
        ],
    )
    _style_sheet(pathways)

    sources = workbook.create_sheet("来源")
    _append_excel_rows(
        sources,
        [
            ["来源名称", "数据集 accession", "版本", "SHA-256"],
            *[
                [
                    source.source_name,
                    source.accession,
                    source.version,
                    source.sha256 or "",
                ]
                for source in result.sources
            ],
        ],
    )
    _style_sheet(sources)
    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def _set_run_font(run, *, size: float = 11, bold: bool = False) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold


def _configure_document(document: DocxDocument) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, color, before, after in (
        ("Heading 1", 16, _BLUE, 18, 10),
        ("Heading 2", 13, _BLUE, 14, 7),
        ("Heading 3", 12, _DARK_BLUE, 10, 5),
    ):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(
            qn("w:eastAsia"),
            "Microsoft YaHei",
        )
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("VetEvidence Agent | 计算预测")
    _set_run_font(footer_run, size=9)
    footer_run.font.color.rgb = RGBColor(89, 89, 89)


def _set_cell_shading(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)


def _set_cell_margins(cell) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, width in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        element = margins.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            margins.append(element)
        element.set(qn("w:w"), str(width))
        element.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != 9360:
        raise ValueError("Word 表格列宽必须合计 9360 DXA。")
    table.autofit = False
    properties = table._tbl.tblPr
    width = properties.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:w"), "9360")
    width.set(qn("w:type"), "dxa")
    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), "0")
    indent.set(qn("w:type"), "dxa")
    grid_columns = table._tbl.tblGrid.gridCol_lst
    for grid_column, column_width in zip(
        grid_columns,
        widths_dxa,
        strict=True,
    ):
        grid_column.set(qn("w:w"), str(column_width))
    for row in table.rows:
        for cell, column_width in zip(row.cells, widths_dxa, strict=True):
            properties = cell._tc.get_or_add_tcPr()
            cell_width = properties.find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                properties.append(cell_width)
            cell_width.set(qn("w:w"), str(column_width))
            cell_width.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _add_docx_table(
    document: DocxDocument,
    headers: list[str],
    rows: list[list[object]],
    widths_dxa: list[int],
) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_row = table.rows[0]
    header_properties = header_row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    header_properties.append(repeat)
    for cell, value in zip(header_row.cells, headers, strict=True):
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(value)
        _set_run_font(run, size=9, bold=True)
        run.font.color.rgb = RGBColor.from_string(_WHITE)
        _set_cell_shading(cell, _BLUE)
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values, strict=True):
            cell.text = ""
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(str(value))
            _set_run_font(run, size=9)
    _set_table_geometry(table, widths_dxa)
    document.add_paragraph()


def network_result_to_docx(result: NetworkPharmacologyResult) -> bytes:
    """Export a compact, traceable Word report for targets and pathways."""

    document = Document()
    _configure_document(document)
    title = (
        document.paragraphs[0]
        if document.paragraphs
        else document.add_paragraph()
    )
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("VetEvidence 网络药理学靶点与通路报告")
    _set_run_font(title_run, size=22, bold=True)
    title_run.font.color.rgb = RGBColor.from_string(_DARK_BLUE)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle_run = subtitle.add_run(
        "可追溯计算预测 | 不等同于文献证据、结合实验或协同证明"
    )
    _set_run_font(subtitle_run, size=10)
    subtitle_run.font.color.rgb = RGBColor(122, 90, 0)

    document.add_heading("摘要", level=1)
    _add_docx_table(
        document,
        ["项目", "值"],
        [
            ["证据等级", result.evidence_grade.value],
            ["算法版本", result.parameters.algorithm_version],
            ["排名规则", result.parameters.ranking_method],
            ["输入化合物数", result.summary.input_compound_count],
            ["交集靶点数", result.summary.intersection_target_count],
            ["交集通路数", result.summary.intersection_pathway_count],
        ],
        [2160, 7200],
    )

    document.add_heading("靶点排名", level=1)
    _add_docx_table(
        document,
        ["排名", "研究对象", "靶点", "靶点 accession", "化合物", "网络分数"],
        [
            [
                target.rank,
                target.organism,
                target.target,
                target.target_accession,
                "；".join(
                    f"{link.compound} ({link.compound_accession})"
                    for link in target.compounds
                ),
                target.network_score,
            ]
            for target in result.ranked_targets
        ],
        [500, 1600, 1400, 1800, 2960, 1100],
    )

    document.add_heading("靶点—通路关系", level=1)
    _add_docx_table(
        document,
        ["靶点", "靶点 accession", "通路", "通路 accession", "来源行"],
        [
            [row[2], row[3], row[4], row[5], row[6]]
            for row in _pathway_rows(result)
        ],
        [1500, 1900, 1900, 2200, 1860],
    )

    document.add_heading("输入来源", level=1)
    _add_docx_table(
        document,
        ["来源名称", "数据集 accession", "版本", "SHA-256"],
        [
            [
                source.source_name,
                source.accession,
                source.version,
                source.sha256 or "",
            ]
            for source in result.sources
        ],
        [1800, 1900, 1200, 4460],
    )
    boundary = document.add_paragraph()
    boundary_run = boundary.add_run(
        "科研边界：网络排名只反映用户导入关系的透明拓扑统计，"
        "不能证明靶点有效、药物结合、体内外活性或药物协同。"
    )
    _set_run_font(boundary_run, size=10, bold=True)
    boundary_run.font.color.rgb = RGBColor(155, 28, 28)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _template_xlsx(columns: tuple[str, ...], *, sheet_name: str) -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = sheet_name
    _append_excel_rows(sheet, [list(columns)])
    _style_sheet(sheet)
    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def _template_docx(columns: tuple[str, ...], *, title_text: str) -> bytes:
    document = Document()
    _configure_document(document)
    title = (
        document.paragraphs[0]
        if document.paragraphs
        else document.add_paragraph()
    )
    title_run = title.add_run(title_text)
    _set_run_font(title_run, size=18, bold=True)
    title_run.font.color.rgb = RGBColor.from_string(_DARK_BLUE)
    note = document.add_paragraph(
        "请保留第一行英文列名，在下方逐行填写静态值；不要合并单元格。"
    )
    note.runs[0].font.color.rgb = RGBColor(89, 89, 89)
    widths = [9360 // len(columns)] * len(columns)
    widths[-1] += 9360 - sum(widths)
    _add_docx_table(
        document,
        list(columns),
        [["" for _ in columns]],
        widths,
    )
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def compound_target_template_xlsx() -> bytes:
    return _template_xlsx(COMPOUND_TARGET_COLUMNS, sheet_name="化合物-靶点")


def target_pathway_template_xlsx() -> bytes:
    return _template_xlsx(TARGET_PATHWAY_COLUMNS, sheet_name="靶点-通路")


def compound_target_template_docx() -> bytes:
    return _template_docx(
        COMPOUND_TARGET_COLUMNS,
        title_text="化合物—靶点关系表模板",
    )


def target_pathway_template_docx() -> bytes:
    return _template_docx(
        TARGET_PATHWAY_COLUMNS,
        title_text="靶点—通路关系表模板",
    )


__all__ = [
    "COMPOUND_TARGET_COLUMNS",
    "MAX_NETWORK_DATA_ROWS",
    "MAX_NETWORK_FILE_BYTES",
    "SUPPORTED_NETWORK_SUFFIXES",
    "TARGET_PATHWAY_COLUMNS",
    "analyze_network_pharmacology_files",
    "compound_target_template_docx",
    "compound_target_template_xlsx",
    "network_result_to_docx",
    "network_result_to_xlsx",
    "parse_compound_target_file",
    "parse_target_pathway_file",
    "target_pathway_template_docx",
    "target_pathway_template_xlsx",
]
